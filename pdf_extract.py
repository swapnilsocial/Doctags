import docx


def docx_tokenize():
    doc = docx.Document('/home/swap/swapnilsocial/Doctags/data/sampledata/sample1.docx')  # Creating word reader object.
    # print(doc.paragraphs)
    text = ''''''
    for para in doc.paragraphs:
        text = text + " " + para.text

    print(text)

docx_tokenize()